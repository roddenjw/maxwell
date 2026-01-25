"""
Tests for export formatting preservation.
Verifies that rich text formatting (bold, italic, underline) is preserved
when exporting from Lexical state to DOCX and PDF formats.
"""

import pytest
import json
from io import BytesIO
from unittest.mock import MagicMock, patch
from docx import Document

from app.services.lexical_utils import (
    TextRun,
    RichParagraph,
    rich_paragraphs_to_lexical_json,
    lexical_to_rich_paragraphs,
    FORMAT_BOLD,
    FORMAT_ITALIC,
    FORMAT_UNDERLINE,
)
from app.services.export_service import ExportService


@pytest.fixture
def mock_db():
    """Create a mock database session."""
    return MagicMock()


@pytest.fixture
def mock_manuscript():
    """Create a mock manuscript."""
    manuscript = MagicMock()
    manuscript.id = "test-manuscript-id"
    manuscript.title = "Test Manuscript"
    manuscript.author = "Test Author"
    return manuscript


@pytest.fixture
def sample_lexical_with_formatting():
    """Create a sample Lexical state with bold, italic, and underline formatting."""
    return json.dumps({
        "root": {
            "children": [
                {
                    "children": [
                        {
                            "detail": 0,
                            "format": 0,  # Plain text
                            "mode": "normal",
                            "style": "",
                            "text": "This is ",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": FORMAT_BOLD,  # Bold
                            "mode": "normal",
                            "style": "",
                            "text": "bold text",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": 0,
                            "mode": "normal",
                            "style": "",
                            "text": " and ",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": FORMAT_ITALIC,  # Italic
                            "mode": "normal",
                            "style": "",
                            "text": "italic text",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": 0,
                            "mode": "normal",
                            "style": "",
                            "text": " and ",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": FORMAT_UNDERLINE,  # Underline
                            "mode": "normal",
                            "style": "",
                            "text": "underlined text",
                            "type": "text",
                            "version": 1
                        },
                        {
                            "detail": 0,
                            "format": 0,
                            "mode": "normal",
                            "style": "",
                            "text": ".",
                            "type": "text",
                            "version": 1
                        }
                    ],
                    "direction": "ltr",
                    "format": "",
                    "indent": 0,
                    "type": "paragraph",
                    "version": 1
                }
            ],
            "direction": "ltr",
            "format": "",
            "indent": 0,
            "type": "root",
            "version": 1
        }
    })


class TestLexicalToRichParagraphs:
    """Test Lexical to RichParagraph conversion."""

    def test_plain_text_conversion(self):
        """Test converting plain text paragraphs."""
        lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": 0,
                                "mode": "normal",
                                "style": "",
                                "text": "Hello, world!",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "paragraph",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })

        paragraphs = lexical_to_rich_paragraphs(lexical_state)

        assert len(paragraphs) == 1
        assert len(paragraphs[0].runs) == 1
        assert paragraphs[0].runs[0].text == "Hello, world!"
        assert paragraphs[0].runs[0].bold is False
        assert paragraphs[0].runs[0].italic is False
        assert paragraphs[0].runs[0].underline is False

    def test_bold_text_conversion(self):
        """Test converting bold text."""
        lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": FORMAT_BOLD,
                                "mode": "normal",
                                "style": "",
                                "text": "Bold text",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "paragraph",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })

        paragraphs = lexical_to_rich_paragraphs(lexical_state)

        assert len(paragraphs) == 1
        assert paragraphs[0].runs[0].bold is True
        assert paragraphs[0].runs[0].italic is False

    def test_mixed_formatting_conversion(self, sample_lexical_with_formatting):
        """Test converting text with mixed formatting."""
        paragraphs = lexical_to_rich_paragraphs(sample_lexical_with_formatting)

        assert len(paragraphs) == 1
        runs = paragraphs[0].runs

        # Find the bold run
        bold_runs = [r for r in runs if r.bold]
        assert len(bold_runs) == 1
        assert bold_runs[0].text == "bold text"

        # Find the italic run
        italic_runs = [r for r in runs if r.italic]
        assert len(italic_runs) == 1
        assert italic_runs[0].text == "italic text"

        # Find the underline run
        underline_runs = [r for r in runs if r.underline]
        assert len(underline_runs) == 1
        assert underline_runs[0].text == "underlined text"

    def test_heading_conversion(self):
        """Test converting heading paragraphs."""
        lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": 0,
                                "mode": "normal",
                                "style": "",
                                "text": "Chapter Title",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "heading",
                        "tag": "h2",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })

        paragraphs = lexical_to_rich_paragraphs(lexical_state)

        assert len(paragraphs) == 1
        assert paragraphs[0].heading_level == 2
        assert paragraphs[0].runs[0].text == "Chapter Title"

    def test_empty_lexical_state(self):
        """Test handling empty Lexical state."""
        paragraphs = lexical_to_rich_paragraphs("")
        assert paragraphs == []

        paragraphs = lexical_to_rich_paragraphs("{}")
        assert paragraphs == []

    def test_invalid_json(self):
        """Test handling invalid JSON."""
        paragraphs = lexical_to_rich_paragraphs("not valid json")
        assert paragraphs == []


class TestRichParagraphsToLexical:
    """Test RichParagraph to Lexical conversion."""

    def test_plain_paragraph_conversion(self):
        """Test converting plain paragraphs to Lexical."""
        paragraphs = [
            RichParagraph(runs=[TextRun(text="Hello, world!")])
        ]

        lexical_json = rich_paragraphs_to_lexical_json(paragraphs)
        state = json.loads(lexical_json)

        assert "root" in state
        assert len(state["root"]["children"]) == 1
        assert state["root"]["children"][0]["type"] == "paragraph"

    def test_formatted_text_conversion(self):
        """Test converting formatted text to Lexical."""
        paragraphs = [
            RichParagraph(runs=[
                TextRun(text="Bold", bold=True),
                TextRun(text=" and "),
                TextRun(text="Italic", italic=True),
            ])
        ]

        lexical_json = rich_paragraphs_to_lexical_json(paragraphs)
        state = json.loads(lexical_json)

        children = state["root"]["children"][0]["children"]
        assert len(children) == 3

        # Check bold
        assert children[0]["format"] == FORMAT_BOLD
        assert children[0]["text"] == "Bold"

        # Check plain
        assert children[1]["format"] == 0
        assert children[1]["text"] == " and "

        # Check italic
        assert children[2]["format"] == FORMAT_ITALIC
        assert children[2]["text"] == "Italic"

    def test_heading_conversion(self):
        """Test converting headings to Lexical."""
        paragraphs = [
            RichParagraph(
                runs=[TextRun(text="Chapter One")],
                heading_level=1
            )
        ]

        lexical_json = rich_paragraphs_to_lexical_json(paragraphs)
        state = json.loads(lexical_json)

        assert state["root"]["children"][0]["type"] == "heading"
        assert state["root"]["children"][0]["tag"] == "h1"


class TestExportServiceFormatting:
    """Test export service formatting preservation."""

    @pytest.mark.asyncio
    async def test_docx_export_preserves_bold(self, mock_db, mock_manuscript):
        """Test that DOCX export preserves bold formatting."""
        chapter = MagicMock()
        chapter.id = "test-chapter-id"
        chapter.title = "Chapter 1"
        chapter.content = "Plain text"
        chapter.lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": FORMAT_BOLD,
                                "mode": "normal",
                                "style": "",
                                "text": "Bold text",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "paragraph",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })
        chapter.word_count = 2

        mock_db.query.return_value.filter_by.return_value.first.return_value = mock_manuscript
        mock_db.query.return_value.filter.return_value.order_by.return_value.all.return_value = [chapter]

        service = ExportService(mock_db)
        result = await service.export_to_docx(mock_manuscript.id)

        # Read the DOCX file and verify bold formatting
        doc = Document(result)

        # Find the paragraph with our bold text (skip title and heading)
        found_bold = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text == "Bold text":
                    assert run.bold is True, "Bold formatting was not preserved"
                    found_bold = True

        assert found_bold, "Bold text was not found in document"

    @pytest.mark.asyncio
    async def test_docx_export_preserves_italic(self, mock_db, mock_manuscript):
        """Test that DOCX export preserves italic formatting."""
        chapter = MagicMock()
        chapter.id = "test-chapter-id"
        chapter.title = "Chapter 1"
        chapter.content = "Plain text"
        chapter.lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": FORMAT_ITALIC,
                                "mode": "normal",
                                "style": "",
                                "text": "Italic text",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "paragraph",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })
        chapter.word_count = 2

        mock_db.query.return_value.filter_by.return_value.first.return_value = mock_manuscript
        mock_db.query.return_value.filter.return_value.order_by.return_value.all.return_value = [chapter]

        service = ExportService(mock_db)
        result = await service.export_to_docx(mock_manuscript.id)

        # Read the DOCX file and verify italic formatting
        doc = Document(result)

        found_italic = False
        for para in doc.paragraphs:
            for run in para.runs:
                if run.text == "Italic text":
                    assert run.italic is True, "Italic formatting was not preserved"
                    found_italic = True

        assert found_italic, "Italic text was not found in document"

    @pytest.mark.asyncio
    async def test_docx_fallback_to_plain_content(self, mock_db, mock_manuscript):
        """Test that DOCX export falls back to plain content when lexical_state is empty."""
        chapter = MagicMock()
        chapter.id = "test-chapter-id"
        chapter.title = "Chapter 1"
        chapter.content = "This is plain text content."
        chapter.lexical_state = ""  # Empty lexical state
        chapter.word_count = 5

        mock_db.query.return_value.filter_by.return_value.first.return_value = mock_manuscript
        mock_db.query.return_value.filter.return_value.order_by.return_value.all.return_value = [chapter]

        service = ExportService(mock_db)
        result = await service.export_to_docx(mock_manuscript.id)

        # Read the DOCX file and verify content exists
        doc = Document(result)

        found_content = False
        for para in doc.paragraphs:
            if "plain text content" in para.text:
                found_content = True

        assert found_content, "Plain text content was not found in document"

    @pytest.mark.asyncio
    async def test_pdf_export_with_formatting(self, mock_db, mock_manuscript):
        """Test that PDF export includes formatting markup."""
        chapter = MagicMock()
        chapter.id = "test-chapter-id"
        chapter.title = "Chapter 1"
        chapter.content = "Plain text"
        chapter.lexical_state = json.dumps({
            "root": {
                "children": [
                    {
                        "children": [
                            {
                                "detail": 0,
                                "format": FORMAT_BOLD,
                                "mode": "normal",
                                "style": "",
                                "text": "Bold",
                                "type": "text",
                                "version": 1
                            },
                            {
                                "detail": 0,
                                "format": FORMAT_ITALIC,
                                "mode": "normal",
                                "style": "",
                                "text": "Italic",
                                "type": "text",
                                "version": 1
                            }
                        ],
                        "direction": "ltr",
                        "format": "",
                        "indent": 0,
                        "type": "paragraph",
                        "version": 1
                    }
                ],
                "direction": "ltr",
                "format": "",
                "indent": 0,
                "type": "root",
                "version": 1
            }
        })
        chapter.word_count = 2

        mock_db.query.return_value.filter_by.return_value.first.return_value = mock_manuscript
        mock_db.query.return_value.filter.return_value.order_by.return_value.all.return_value = [chapter]

        service = ExportService(mock_db)

        # PDF export should complete without error
        result = await service.export_to_pdf(mock_manuscript.id)

        # Verify we got a BytesIO object with content
        assert isinstance(result, BytesIO)
        content = result.getvalue()
        assert len(content) > 0
        # PDF files start with %PDF
        assert content[:4] == b'%PDF'

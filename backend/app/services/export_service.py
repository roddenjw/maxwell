"""
Export Service
Handles manuscript export to various formats (DOCX, PDF, EPUB)
"""

from io import BytesIO
from typing import Optional, List
from sqlalchemy.orm import Session
from sqlalchemy import and_

from app.models.manuscript import Manuscript, Chapter
from app.services.lexical_utils import lexical_to_rich_paragraphs, RichParagraph


class ExportService:
    """Service for exporting manuscripts to different formats"""

    def __init__(self, db: Session):
        self.db = db

    async def export_to_docx(
        self,
        manuscript_id: str,
        include_folders: bool = False,
        chapter_ids: Optional[List[str]] = None
    ) -> BytesIO:
        """
        Export manuscript to DOCX format

        Args:
            manuscript_id: ID of the manuscript to export
            include_folders: Whether to include folder names as headings
            chapter_ids: Optional list of specific chapter IDs to export

        Returns:
            BytesIO buffer containing the DOCX file
        """
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        # Create document
        doc = Document()

        # Set page margins (1 inch on all sides - standard manuscript format)
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Get manuscript
        manuscript = self.db.query(Manuscript).filter_by(id=manuscript_id).first()
        if not manuscript:
            raise ValueError(f"Manuscript {manuscript_id} not found")

        # Add title page
        title_paragraph = doc.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_paragraph.add_run(manuscript.title)
        title_run.font.size = Pt(24)
        title_run.font.bold = True

        if manuscript.author:
            author_paragraph = doc.add_paragraph()
            author_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            author_run = author_paragraph.add_run(f"by {manuscript.author}")
            author_run.font.size = Pt(14)

        # Add page break after title
        doc.add_page_break()

        # Get chapters
        if chapter_ids:
            # Export specific chapters
            chapters = []
            for chapter_id in chapter_ids:
                chapter = self.db.query(Chapter).filter_by(id=chapter_id).first()
                if chapter:
                    chapters.append(chapter)
        else:
            # Export all chapters in order
            chapters = self.db.query(Chapter).filter(
                and_(
                    Chapter.manuscript_id == manuscript_id,
                    Chapter.is_folder == False
                )
            ).order_by(Chapter.order_index).all()

        # Add chapters
        for chapter in chapters:
            # Add chapter title as heading
            heading = doc.add_heading(chapter.title, level=1)
            heading_run = heading.runs[0]
            heading_run.font.size = Pt(18)

            # Add chapter content with formatting preserved
            rich_paragraphs = []
            if chapter.lexical_state and chapter.lexical_state.strip():
                # Parse Lexical state to extract formatting
                rich_paragraphs = lexical_to_rich_paragraphs(chapter.lexical_state)

            if rich_paragraphs:
                # Write formatted paragraphs
                for rich_para in rich_paragraphs:
                    if rich_para.is_empty():
                        continue

                    # Handle headings vs regular paragraphs
                    if rich_para.heading_level > 0:
                        para = doc.add_heading('', level=min(rich_para.heading_level, 6))
                    else:
                        para = doc.add_paragraph()
                        # Standard manuscript formatting for body text
                        para_format = para.paragraph_format
                        para_format.line_spacing = 2.0  # Double-spaced
                        para_format.first_line_indent = Inches(0.5)  # Indent first line

                    # Add runs with formatting
                    for text_run in rich_para.runs:
                        if not text_run.text:
                            continue
                        run = para.add_run(text_run.text)
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)
                        run.bold = text_run.bold
                        run.italic = text_run.italic
                        run.underline = text_run.underline
                        if text_run.strikethrough:
                            run.font.strike = True

            elif chapter.content:
                # Fallback: Split plain text into paragraphs
                paragraphs = chapter.content.split('\n')

                for para_text in paragraphs:
                    if para_text.strip():
                        para = doc.add_paragraph(para_text.strip())
                        # Standard manuscript formatting
                        para_format = para.paragraph_format
                        para_format.line_spacing = 2.0  # Double-spaced
                        para_format.first_line_indent = Inches(0.5)  # Indent first line

                        # Set font
                        for run in para.runs:
                            run.font.name = 'Times New Roman'
                            run.font.size = Pt(12)

            # Page break between chapters
            doc.add_page_break()

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer

    async def export_to_pdf(
        self,
        manuscript_id: str,
        include_folders: bool = False,
        chapter_ids: Optional[List[str]] = None
    ) -> BytesIO:
        """
        Export manuscript to PDF format

        Strategy: Convert to DOCX first, then to PDF using docx2pdf or reportlab

        Args:
            manuscript_id: ID of the manuscript to export
            include_folders: Whether to include folder names as headings
            chapter_ids: Optional list of specific chapter IDs to export

        Returns:
            BytesIO buffer containing the PDF file
        """
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY

        # Create buffer
        buffer = BytesIO()

        # Create PDF document
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            topMargin=1*inch,
            bottomMargin=1*inch,
            leftMargin=1*inch,
            rightMargin=1*inch
        )

        # Get manuscript
        manuscript = self.db.query(Manuscript).filter_by(id=manuscript_id).first()
        if not manuscript:
            raise ValueError(f"Manuscript {manuscript_id} not found")

        # Container for the 'Flowable' objects
        elements = []

        # Define styles
        styles = getSampleStyleSheet()

        # Title style
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Title'],
            fontSize=24,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        # Author style
        author_style = ParagraphStyle(
            'Author',
            parent=styles['Normal'],
            fontSize=14,
            textColor='black',
            spaceAfter=30,
            alignment=TA_CENTER
        )

        # Chapter heading style
        chapter_style = ParagraphStyle(
            'ChapterHeading',
            parent=styles['Heading1'],
            fontSize=18,
            textColor='black',
            spaceAfter=20,
            spaceBefore=20
        )

        # Body text style (manuscript format)
        body_style = ParagraphStyle(
            'ManuscriptBody',
            parent=styles['Normal'],
            fontSize=12,
            leading=24,  # Double spacing
            firstLineIndent=0.5*inch,
            alignment=TA_JUSTIFY,
            fontName='Times-Roman'
        )

        # Add title page
        elements.append(Paragraph(manuscript.title, title_style))
        if manuscript.author:
            elements.append(Paragraph(f"by {manuscript.author}", author_style))

        elements.append(PageBreak())

        # Get chapters
        if chapter_ids:
            chapters = []
            for chapter_id in chapter_ids:
                chapter = self.db.query(Chapter).filter_by(id=chapter_id).first()
                if chapter:
                    chapters.append(chapter)
        else:
            chapters = self.db.query(Chapter).filter(
                and_(
                    Chapter.manuscript_id == manuscript_id,
                    Chapter.is_folder == False
                )
            ).order_by(Chapter.order_index).all()

        def escape_for_reportlab(text: str) -> str:
            """Escape special characters for ReportLab XML."""
            return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

        def format_run_for_pdf(run) -> str:
            """Convert a TextRun to ReportLab markup."""
            text = escape_for_reportlab(run.text)
            if run.bold:
                text = f'<b>{text}</b>'
            if run.italic:
                text = f'<i>{text}</i>'
            if run.underline:
                text = f'<u>{text}</u>'
            if run.strikethrough:
                text = f'<strike>{text}</strike>'
            return text

        # Add chapters
        for i, chapter in enumerate(chapters):
            # Add chapter title
            elements.append(Paragraph(chapter.title, chapter_style))
            elements.append(Spacer(1, 0.2*inch))

            # Add chapter content with formatting preserved
            rich_paragraphs = []
            if chapter.lexical_state and chapter.lexical_state.strip():
                rich_paragraphs = lexical_to_rich_paragraphs(chapter.lexical_state)

            if rich_paragraphs:
                for rich_para in rich_paragraphs:
                    if rich_para.is_empty():
                        continue

                    # Build paragraph text with formatting tags
                    para_text = ''.join(format_run_for_pdf(run) for run in rich_para.runs if run.text)

                    if para_text.strip():
                        # Use heading style for headings
                        if rich_para.heading_level > 0:
                            elements.append(Paragraph(para_text, chapter_style))
                        else:
                            elements.append(Paragraph(para_text, body_style))
                        elements.append(Spacer(1, 0.1*inch))

            elif chapter.content:
                # Fallback: Split plain text into paragraphs
                paragraphs = chapter.content.split('\n')

                for para_text in paragraphs:
                    if para_text.strip():
                        safe_text = escape_for_reportlab(para_text.strip())
                        elements.append(Paragraph(safe_text, body_style))
                        elements.append(Spacer(1, 0.1*inch))

            # Page break between chapters (except last)
            if i < len(chapters) - 1:
                elements.append(PageBreak())

        # Build PDF
        doc.build(elements)
        buffer.seek(0)

        return buffer

    async def get_export_preview(self, manuscript_id: str) -> dict:
        """
        Get preview information for export

        Args:
            manuscript_id: ID of the manuscript

        Returns:
            Dictionary with export preview data
        """
        manuscript = self.db.query(Manuscript).filter_by(id=manuscript_id).first()
        if not manuscript:
            raise ValueError(f"Manuscript {manuscript_id} not found")

        chapters = self.db.query(Chapter).filter(
            and_(
                Chapter.manuscript_id == manuscript_id,
                Chapter.is_folder == False
            )
        ).order_by(Chapter.order_index).all()

        total_words = sum(chapter.word_count or 0 for chapter in chapters)

        return {
            "manuscript_id": manuscript_id,
            "title": manuscript.title,
            "author": manuscript.author,
            "total_chapters": len(chapters),
            "total_words": total_words,
            "chapters": [
                {
                    "id": chapter.id,
                    "title": chapter.title,
                    "word_count": chapter.word_count or 0,
                    "order_index": chapter.order_index
                }
                for chapter in chapters
            ]
        }

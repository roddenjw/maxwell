"""
Import Service
Handles manuscript import from various document formats (DOCX, RTF, ODT, TXT, MD, PDF)
"""

import re
import uuid
import json
from io import BytesIO
from dataclasses import dataclass, field
from typing import List, Optional, Tuple
from datetime import datetime, timedelta

from sqlalchemy.orm import Session

from app.models.manuscript import Manuscript, Chapter
from app.services.lexical_utils import (
    TextRun,
    RichParagraph,
    rich_paragraphs_to_lexical_json,
    FORMAT_BOLD,
    FORMAT_ITALIC,
    FORMAT_UNDERLINE,
)


# Supported formats with metadata
SUPPORTED_FORMATS = {
    ".docx": {
        "extension": ".docx",
        "name": "Microsoft Word",
        "description": "Word documents (.docx)",
        "formatting_support": "full",  # full, partial, none
        "mime_types": [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ],
    },
    ".rtf": {
        "extension": ".rtf",
        "name": "Rich Text Format",
        "description": "Rich Text Format files (.rtf)",
        "formatting_support": "none",  # striprtf only extracts plain text
        "mime_types": ["application/rtf", "text/rtf"],
    },
    ".odt": {
        "extension": ".odt",
        "name": "OpenDocument Text",
        "description": "OpenDocument text files (.odt)",
        "formatting_support": "full",
        "mime_types": ["application/vnd.oasis.opendocument.text"],
    },
    ".txt": {
        "extension": ".txt",
        "name": "Plain Text",
        "description": "Plain text files (.txt)",
        "formatting_support": "none",
        "mime_types": ["text/plain"],
    },
    ".md": {
        "extension": ".md",
        "name": "Markdown",
        "description": "Markdown files (.md)",
        "formatting_support": "partial",
        "mime_types": ["text/markdown", "text/x-markdown"],
    },
    ".pdf": {
        "extension": ".pdf",
        "name": "PDF",
        "description": "PDF files (.pdf) - best effort extraction",
        "formatting_support": "partial",
        "mime_types": ["application/pdf"],
        "warning": "PDF import is lossy. Formatting may not be perfectly preserved.",
    },
}


@dataclass
class ChapterBoundary:
    """Represents a detected chapter boundary in the document."""
    start_index: int  # Index in the paragraph list
    end_index: int  # Exclusive end index
    title: str
    detection_method: str  # "heading", "pattern", "page_break", "single"


@dataclass
class ParsedDocument:
    """Result of parsing a document file."""
    paragraphs: List[RichParagraph] = field(default_factory=list)
    title: Optional[str] = None
    author: Optional[str] = None
    page_breaks: List[int] = field(default_factory=list)  # Paragraph indices after page breaks
    warnings: List[str] = field(default_factory=list)


@dataclass
class DetectedChapter:
    """A chapter detected from the imported document."""
    index: int
    title: str
    paragraphs: List[RichParagraph]
    lexical_state: str  # Pre-computed Lexical JSON
    plain_content: str  # Plain text for word count
    word_count: int


@dataclass
class ImportResult:
    """Result of parsing and detecting chapters in a document."""
    parse_id: str
    title: str
    author: Optional[str]
    total_words: int
    detection_method: str
    format_warnings: List[str]
    chapters: List[DetectedChapter]
    source_format: str


# In-memory cache for parsed documents (30-minute TTL)
_parse_cache: dict[str, Tuple[ImportResult, datetime]] = {}
CACHE_TTL_MINUTES = 30


def _clean_cache():
    """Remove expired cache entries."""
    now = datetime.utcnow()
    expired_keys = [
        key for key, (_, created_at) in _parse_cache.items()
        if now - created_at > timedelta(minutes=CACHE_TTL_MINUTES)
    ]
    for key in expired_keys:
        del _parse_cache[key]


class ImportService:
    """Service for importing documents into Maxwell."""

    def __init__(self):
        pass

    @staticmethod
    def get_supported_formats() -> List[dict]:
        """Get list of supported import formats with metadata."""
        return list(SUPPORTED_FORMATS.values())

    async def parse_document(
        self,
        file_content: bytes,
        filename: str,
        detection_mode: str = "auto"
    ) -> ImportResult:
        """
        Parse a document and detect chapters.

        Args:
            file_content: Raw file bytes
            filename: Original filename (used to detect format)
            detection_mode: How to detect chapters - "auto", "headings", "pattern", "page_breaks", "single"

        Returns:
            ImportResult with parsed chapters and metadata
        """
        # Clean expired cache entries
        _clean_cache()

        # Determine format from extension
        ext = self._get_extension(filename)
        if ext not in SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {ext}")

        # Parse the document
        parsed = await self._parse_by_format(file_content, ext)

        # Detect chapters
        chapters = self._detect_chapters(parsed, detection_mode)

        # Extract title from filename if not found in document
        title = parsed.title or self._title_from_filename(filename)

        # Calculate total word count
        total_words = sum(ch.word_count for ch in chapters)

        # Build result
        result = ImportResult(
            parse_id=str(uuid.uuid4()),
            title=title,
            author=parsed.author,
            total_words=total_words,
            detection_method=detection_mode if detection_mode != "auto" else self._get_detection_method(chapters),
            format_warnings=parsed.warnings,
            chapters=chapters,
            source_format=ext,
        )

        # Cache the result
        _parse_cache[result.parse_id] = (result, datetime.utcnow())

        return result

    async def create_manuscript_from_import(
        self,
        db: Session,
        parse_id: str,
        title: Optional[str] = None,
        author: Optional[str] = None,
        description: Optional[str] = None,
        chapter_adjustments: Optional[List[dict]] = None,
        series_id: Optional[str] = None,
    ) -> Manuscript:
        """
        Create a manuscript from a previously parsed import.

        Args:
            db: Database session
            parse_id: ID from parse_document result
            title: Override title (optional)
            author: Override author (optional)
            description: Manuscript description (optional)
            chapter_adjustments: List of chapter modifications (optional)
                Each dict: { "index": int, "title": str (optional), "included": bool (optional) }
            series_id: ID of series to add manuscript to (optional)

        Returns:
            Created Manuscript with chapters
        """
        # Get cached parse result
        if parse_id not in _parse_cache:
            raise ValueError("Parse result not found or expired. Please re-upload the file.")

        result, _ = _parse_cache.pop(parse_id)

        # Apply chapter adjustments
        included_chapters = []
        for ch in result.chapters:
            # Check if this chapter was excluded or renamed
            if chapter_adjustments:
                adjustment = next(
                    (adj for adj in chapter_adjustments if adj.get("index") == ch.index),
                    None
                )
                if adjustment:
                    if adjustment.get("included") is False:
                        continue  # Skip excluded chapters
                    if adjustment.get("title"):
                        ch.title = adjustment["title"]

            included_chapters.append(ch)

        # Create manuscript
        manuscript = Manuscript(
            id=str(uuid.uuid4()),
            title=title or result.title,
            author=author or result.author or "",
            description=description or f"Imported from {result.source_format} file",
            word_count=sum(ch.word_count for ch in included_chapters),
            series_id=series_id,
        )
        db.add(manuscript)
        db.flush()  # Get the ID

        # Create chapters
        for order_idx, ch in enumerate(included_chapters):
            chapter = Chapter(
                id=str(uuid.uuid4()),
                manuscript_id=manuscript.id,
                title=ch.title,
                is_folder=0,
                order_index=order_idx,
                lexical_state=ch.lexical_state,
                content=ch.plain_content,
                word_count=ch.word_count,
            )
            db.add(chapter)

        db.commit()
        db.refresh(manuscript)

        return manuscript

    def get_cached_result(self, parse_id: str) -> Optional[ImportResult]:
        """Get a cached parse result by ID."""
        _clean_cache()
        if parse_id in _parse_cache:
            return _parse_cache[parse_id][0]
        return None

    # --- Format Parsers ---

    async def _parse_by_format(self, content: bytes, ext: str) -> ParsedDocument:
        """Route to the appropriate parser based on extension."""
        parsers = {
            ".docx": self._parse_docx,
            ".rtf": self._parse_rtf,
            ".odt": self._parse_odt,
            ".txt": self._parse_txt,
            ".md": self._parse_markdown,
            ".pdf": self._parse_pdf,
        }
        parser = parsers.get(ext)
        if not parser:
            raise ValueError(f"No parser for format: {ext}")
        return await parser(content)

    async def _parse_docx(self, content: bytes) -> ParsedDocument:
        """Parse a DOCX file with formatting preserved."""
        from docx import Document
        from docx.enum.text import WD_BREAK

        doc = Document(BytesIO(content))
        paragraphs = []
        page_breaks = []
        title = None
        author = None

        # Try to extract title/author from core properties
        try:
            if doc.core_properties.title:
                title = doc.core_properties.title
            if doc.core_properties.author:
                author = doc.core_properties.author
        except Exception:
            pass

        for para in doc.paragraphs:
            # Check for page breaks
            for run in para.runs:
                if run._element.xml and 'w:br' in run._element.xml and 'w:type="page"' in run._element.xml:
                    page_breaks.append(len(paragraphs))
                    break

            # Determine heading level
            heading_level = 0
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    heading_level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    heading_level = 1

            # Extract runs with formatting
            runs = []
            for run in para.runs:
                if run.text:
                    runs.append(TextRun(
                        text=run.text,
                        bold=run.bold or False,
                        italic=run.italic or False,
                        underline=run.underline or False,
                    ))

            if runs:  # Only add non-empty paragraphs
                paragraphs.append(RichParagraph(runs=runs, heading_level=heading_level))

        return ParsedDocument(
            paragraphs=paragraphs,
            title=title,
            author=author,
            page_breaks=page_breaks,
        )

    async def _parse_rtf(self, content: bytes) -> ParsedDocument:
        """Parse an RTF file (plain text only)."""
        from striprtf.striprtf import rtf_to_text

        try:
            # Decode RTF content
            text = content.decode('utf-8', errors='replace')
            plain_text = rtf_to_text(text)
        except Exception as e:
            raise ValueError(f"Failed to parse RTF file: {e}")

        # Split into paragraphs
        paragraphs = []
        for line in plain_text.split('\n'):
            if line.strip():
                paragraphs.append(RichParagraph(
                    runs=[TextRun(text=line.strip())],
                    heading_level=0,
                ))

        return ParsedDocument(
            paragraphs=paragraphs,
            warnings=["RTF import extracts plain text only. Formatting was not preserved."],
        )

    async def _parse_odt(self, content: bytes) -> ParsedDocument:
        """Parse an ODT file with formatting preserved."""
        from odf.opendocument import load
        from odf.text import P, H, Span
        from odf import text as odf_text

        try:
            doc = load(BytesIO(content))
        except Exception as e:
            raise ValueError(f"Failed to parse ODT file: {e}")

        paragraphs = []
        title = None
        author = None

        # Try to get metadata
        try:
            meta = doc.meta
            if meta:
                for child in meta.childNodes:
                    if child.tagName == "dc:title" and child.firstChild:
                        title = str(child.firstChild)
                    elif child.tagName == "dc:creator" and child.firstChild:
                        author = str(child.firstChild)
        except Exception:
            pass

        def extract_text_with_formatting(element, default_bold=False, default_italic=False):
            """Recursively extract text and formatting from ODT elements."""
            runs = []

            for node in element.childNodes:
                if node.nodeType == node.TEXT_NODE:
                    text = str(node)
                    if text:
                        runs.append(TextRun(
                            text=text,
                            bold=default_bold,
                            italic=default_italic,
                        ))
                elif hasattr(node, 'tagName'):
                    if node.tagName == "text:span":
                        # Check style for formatting
                        style_name = node.getAttribute("stylename") or ""
                        is_bold = default_bold or "bold" in style_name.lower()
                        is_italic = default_italic or "italic" in style_name.lower()
                        runs.extend(extract_text_with_formatting(node, is_bold, is_italic))
                    else:
                        runs.extend(extract_text_with_formatting(node, default_bold, default_italic))

            return runs

        # Process document body
        for elem in doc.text.childNodes:
            if hasattr(elem, 'tagName'):
                if elem.tagName == "text:p":
                    runs = extract_text_with_formatting(elem)
                    if runs:
                        paragraphs.append(RichParagraph(runs=runs, heading_level=0))
                elif elem.tagName == "text:h":
                    # Heading
                    level = 1
                    outline_level = elem.getAttribute("outlinelevel")
                    if outline_level:
                        try:
                            level = int(outline_level)
                        except ValueError:
                            pass
                    runs = extract_text_with_formatting(elem)
                    if runs:
                        paragraphs.append(RichParagraph(runs=runs, heading_level=level))

        return ParsedDocument(
            paragraphs=paragraphs,
            title=title,
            author=author,
        )

    async def _parse_txt(self, content: bytes) -> ParsedDocument:
        """Parse a plain text file."""
        # Try different encodings - UTF-16 only with BOM to avoid false positives
        text = None

        # Check for UTF-16 BOM first
        if content.startswith(b'\xff\xfe') or content.startswith(b'\xfe\xff'):
            try:
                text = content.decode('utf-16')
            except UnicodeDecodeError:
                pass

        # Try other encodings if UTF-16 didn't work
        if text is None:
            for encoding in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    text = content.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

        if text is None:
            text = content.decode('utf-8', errors='replace')

        paragraphs = []
        for line in text.split('\n'):
            if line.strip():
                paragraphs.append(RichParagraph(
                    runs=[TextRun(text=line.strip())],
                    heading_level=0,
                ))

        return ParsedDocument(paragraphs=paragraphs)

    async def _parse_markdown(self, content: bytes) -> ParsedDocument:
        """Parse a Markdown file with partial formatting support."""
        from markdown_it import MarkdownIt

        text = content.decode('utf-8', errors='replace')
        md = MarkdownIt()
        tokens = md.parse(text)

        paragraphs = []
        current_runs = []
        current_heading_level = 0

        def process_inline(inline_tokens, bold=False, italic=False):
            """Process inline tokens recursively."""
            runs = []
            b, i = bold, italic

            for token in inline_tokens:
                if token.type == 'text':
                    if token.content:
                        runs.append(TextRun(text=token.content, bold=b, italic=i))
                elif token.type == 'strong_open':
                    b = True
                elif token.type == 'strong_close':
                    b = bold  # Reset to parent state
                elif token.type == 'em_open':
                    i = True
                elif token.type == 'em_close':
                    i = italic  # Reset to parent state
                elif token.type == 'softbreak':
                    runs.append(TextRun(text=' ', bold=b, italic=i))
                elif token.children:
                    runs.extend(process_inline(token.children, b, i))

            return runs

        i = 0
        while i < len(tokens):
            token = tokens[i]

            if token.type == 'heading_open':
                # Get heading level from tag (h1, h2, etc.)
                current_heading_level = int(token.tag[1]) if token.tag and len(token.tag) == 2 else 1
                i += 1
            elif token.type == 'heading_close':
                if current_runs:
                    paragraphs.append(RichParagraph(runs=current_runs, heading_level=current_heading_level))
                    current_runs = []
                current_heading_level = 0
                i += 1
            elif token.type == 'paragraph_open':
                i += 1
            elif token.type == 'paragraph_close':
                if current_runs:
                    paragraphs.append(RichParagraph(runs=current_runs, heading_level=current_heading_level))
                    current_runs = []
                i += 1
            elif token.type == 'inline' and token.children:
                current_runs.extend(process_inline(token.children))
                i += 1
            else:
                i += 1

        # Handle any remaining runs
        if current_runs:
            paragraphs.append(RichParagraph(runs=current_runs, heading_level=0))

        return ParsedDocument(
            paragraphs=paragraphs,
            warnings=["Markdown formatting support is partial. Complex elements may not be preserved."],
        )

    async def _parse_pdf(self, content: bytes) -> ParsedDocument:
        """Parse a PDF file (best-effort text extraction)."""
        try:
            import fitz  # pymupdf
        except ImportError:
            raise ValueError("PDF parsing requires pymupdf. Install with: pip install pymupdf")

        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception as e:
            raise ValueError(f"Failed to open PDF file: {e}")

        paragraphs = []
        page_breaks = []
        warnings = [
            "PDF import is lossy. Text extraction is best-effort.",
            "Formatting (bold, italic) is extracted when font metadata is available.",
        ]

        for page_num in range(len(doc)):
            page = doc[page_num]

            # Mark page break (except for first page)
            if page_num > 0:
                page_breaks.append(len(paragraphs))

            # Extract text blocks with formatting
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block.get("type") == 0:  # Text block
                    for line in block.get("lines", []):
                        runs = []
                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text.strip():
                                continue

                            # Try to detect formatting from font
                            font = span.get("font", "").lower()
                            flags = span.get("flags", 0)

                            is_bold = "bold" in font or (flags & 2 ** 4) != 0
                            is_italic = "italic" in font or "oblique" in font or (flags & 2 ** 1) != 0

                            runs.append(TextRun(
                                text=text,
                                bold=is_bold,
                                italic=is_italic,
                            ))

                        if runs:
                            paragraphs.append(RichParagraph(runs=runs, heading_level=0))

        doc.close()

        return ParsedDocument(
            paragraphs=paragraphs,
            page_breaks=page_breaks,
            warnings=warnings,
        )

    # --- Chapter Detection ---

    def _detect_chapters(
        self,
        parsed: ParsedDocument,
        mode: str = "auto"
    ) -> List[DetectedChapter]:
        """
        Detect chapters in a parsed document.

        Args:
            parsed: ParsedDocument with paragraphs
            mode: Detection mode - "auto", "headings", "pattern", "page_breaks", "single"

        Returns:
            List of DetectedChapter objects
        """
        if not parsed.paragraphs:
            return []

        if mode == "single":
            # Treat entire document as one chapter
            return [self._create_chapter(0, "Chapter 1", parsed.paragraphs)]

        boundaries = []

        if mode == "auto":
            # Try detection methods in order of preference
            boundaries = self._detect_by_headings(parsed)
            if not boundaries:
                boundaries = self._detect_by_pattern(parsed)
            if not boundaries:
                boundaries = self._detect_by_page_breaks(parsed)
            if not boundaries:
                # Fall back to single chapter
                return [self._create_chapter(0, "Chapter 1", parsed.paragraphs)]
        elif mode == "headings":
            boundaries = self._detect_by_headings(parsed)
        elif mode == "pattern":
            boundaries = self._detect_by_pattern(parsed)
        elif mode == "page_breaks":
            boundaries = self._detect_by_page_breaks(parsed)

        if not boundaries:
            return [self._create_chapter(0, "Chapter 1", parsed.paragraphs)]

        # Convert boundaries to chapters
        chapters = []
        for i, boundary in enumerate(boundaries):
            paras = parsed.paragraphs[boundary.start_index:boundary.end_index]
            chapters.append(self._create_chapter(i, boundary.title, paras))

        return chapters

    def _detect_by_headings(self, parsed: ParsedDocument) -> List[ChapterBoundary]:
        """Detect chapters by heading styles (H1, H2)."""
        boundaries = []
        current_start = 0

        for i, para in enumerate(parsed.paragraphs):
            if para.heading_level in (1, 2):
                # Close previous chapter if we have content
                if i > current_start:
                    if boundaries:
                        boundaries[-1] = ChapterBoundary(
                            start_index=boundaries[-1].start_index,
                            end_index=i,
                            title=boundaries[-1].title,
                            detection_method="heading",
                        )

                # Start new chapter
                title = para.get_plain_text().strip()
                boundaries.append(ChapterBoundary(
                    start_index=i,
                    end_index=len(parsed.paragraphs),  # Will be updated
                    title=title or f"Chapter {len(boundaries) + 1}",
                    detection_method="heading",
                ))
                current_start = i

        return boundaries

    def _detect_by_pattern(self, parsed: ParsedDocument) -> List[ChapterBoundary]:
        """Detect chapters by text patterns (Chapter 1, CHAPTER ONE, Part I, etc.)."""
        patterns = [
            r'^(?:Chapter|CHAPTER)\s+(\d+|[IVXLCDM]+|[A-Z][a-z]+)',  # Chapter 1, CHAPTER I, Chapter One
            r'^(?:Part|PART)\s+(\d+|[IVXLCDM]+|[A-Z][a-z]+)',  # Part 1, PART I
            r'^(\d+)\.\s+\w',  # "1. Title" format
            r'^(?:Prologue|PROLOGUE|Epilogue|EPILOGUE)$',  # Prologue/Epilogue
        ]
        combined_pattern = '|'.join(f'({p})' for p in patterns)

        boundaries = []

        for i, para in enumerate(parsed.paragraphs):
            text = para.get_plain_text().strip()
            if re.match(combined_pattern, text, re.IGNORECASE):
                # Close previous chapter
                if boundaries:
                    boundaries[-1] = ChapterBoundary(
                        start_index=boundaries[-1].start_index,
                        end_index=i,
                        title=boundaries[-1].title,
                        detection_method="pattern",
                    )

                # Start new chapter
                boundaries.append(ChapterBoundary(
                    start_index=i,
                    end_index=len(parsed.paragraphs),
                    title=text[:50] if len(text) <= 50 else text[:47] + "...",
                    detection_method="pattern",
                ))

        return boundaries

    def _detect_by_page_breaks(self, parsed: ParsedDocument) -> List[ChapterBoundary]:
        """Detect chapters by page breaks."""
        if not parsed.page_breaks:
            return []

        boundaries = []
        prev_start = 0

        for i, break_idx in enumerate(parsed.page_breaks):
            if break_idx > prev_start:
                boundaries.append(ChapterBoundary(
                    start_index=prev_start,
                    end_index=break_idx,
                    title=f"Chapter {len(boundaries) + 1}",
                    detection_method="page_break",
                ))
                prev_start = break_idx

        # Add final chapter
        if prev_start < len(parsed.paragraphs):
            boundaries.append(ChapterBoundary(
                start_index=prev_start,
                end_index=len(parsed.paragraphs),
                title=f"Chapter {len(boundaries) + 1}",
                detection_method="page_break",
            ))

        return boundaries

    def _create_chapter(
        self,
        index: int,
        title: str,
        paragraphs: List[RichParagraph]
    ) -> DetectedChapter:
        """Create a DetectedChapter from paragraphs."""
        # Generate Lexical state
        lexical_state = rich_paragraphs_to_lexical_json(paragraphs)

        # Generate plain text
        plain_text = '\n'.join(p.get_plain_text() for p in paragraphs)

        # Count words
        word_count = len(plain_text.split())

        # Generate preview text (first 200 chars)
        preview = plain_text[:200].strip()
        if len(plain_text) > 200:
            preview += "..."

        return DetectedChapter(
            index=index,
            title=title,
            paragraphs=paragraphs,
            lexical_state=lexical_state,
            plain_content=plain_text,
            word_count=word_count,
        )

    def _get_detection_method(self, chapters: List[DetectedChapter]) -> str:
        """Get the detection method used for chapters."""
        if not chapters:
            return "single"
        # The method is stored in the first paragraph's heading level or detected pattern
        # For simplicity, return "auto" as the auto-detection was used
        return "auto"

    # --- Utilities ---

    def _get_extension(self, filename: str) -> str:
        """Get lowercase file extension from filename."""
        if '.' not in filename:
            raise ValueError("File has no extension")
        return '.' + filename.rsplit('.', 1)[1].lower()

    def _title_from_filename(self, filename: str) -> str:
        """Extract a title from the filename."""
        # Remove extension
        name = filename.rsplit('.', 1)[0]
        # Replace underscores and hyphens with spaces
        name = name.replace('_', ' ').replace('-', ' ')
        # Title case
        return name.title()


# Singleton instance
import_service = ImportService()
